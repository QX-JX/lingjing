import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches

def ensure_dir(path: str):
    os.makedirs(os.path.dirname(path), exist_ok=True)

def create_software_docx(target_path: str, pages: int = 60):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Set default font to a common Chinese font
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # Title page
    title = doc.add_paragraph()
    run = title.add_run('软著材料（样本）')
    run.font.size = Pt(22)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('版本：样本 1.0').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('仅用于版式占位，无具体内容要求').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # Generate content pages
    for i in range(1, pages + 1):
        heading = doc.add_paragraph(f'第 {i} 页')
        heading.runs[0].font.size = Pt(16)

        # Add some filler lines to better occupy the page
        for line in range(1, 30):
            doc.add_paragraph(f'占位内容 第{i}页 - 行{line}')

        # Add manual page break except for last page
        if i < pages:
            doc.add_page_break()

    ensure_dir(target_path)
    doc.save(target_path)

if __name__ == '__main__':
    out_path = os.path.abspath(os.path.join(os.getcwd(), '软著_样本_60页.docx'))
    create_software_docx(out_path, pages=60)
    print(out_path)
